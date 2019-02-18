import datetime
from docx import Document
from docx.shared import Inches
import re
import io
from googleapiclient.http import MediaIoBaseDownload
import os

PROG_PATH = os.path.dirname(__file__)

class MicrosoftDocxParser:
    """ Class that parses a Microsoft Docx format report template.
    """
    
    def __inti__(self):
        pass
    
    def download_report_template(self, drive_service, name, save_path, team_drive_id):
        """ Downloads the report template from google drive
        
        Args:
            service (google drive service): the google drive api service
            name (str): The name of the report template file
            save_path(str): The local path to save the report template file
            team_drive_id (str): the id of the team drive if using Team Drives
            
        Returns:
            bool: True if download was successful
        
        """
        
        #page token is google api way to paginate results
        page_token = None
        while True:
            
            #different calls are required depending on if team drives are being used
            #search for the filename on Google Drive
            if team_drive_id is None:
                response = drive_service.files().list(q="name = '" + name + "' and mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'",
                                                      spaces='drive',
                                                      fields='nextPageToken, files(id, name)',
                                                      pageToken=page_token).execute()
            else:
                response = drive_service.files().list(q="name = '" + name + "' and mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'",
                                                      spaces='drive',
                                                      corpora='teamDrive',
                                                      supportsTeamDrives=True,
                                                      includeTeamDriveItems = True,
                                                      teamDriveId=team_drive_id,
                                                      fields='nextPageToken, files(id, name)',
                                                      pageToken=page_token).execute()
            
            #download the file
            for file in response.get('files', []):
            
                file_id = file.get('id')
                request = drive_service.files().get_media(fileId=file_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
       
                #save the file to local storage
                fo = open(save_path, "wb")
                fo.write(fh.getvalue())
                fo.close()
        
                return True
            
            page_token = response.get('nextPageToken', None)
            if page_token is None:
                break
            
        return False
    
    def extract_table_commands(self, document_path):
        """ Extracts the doc template requests for data in the form of command strings
        
        Args:
            document_path (str): the path of the template doc file
            
        Returns:
            str[]: Array of table commands extracted from the template
        
        """
        
        document = Document(document_path)
        
        #go through the table objects and extract the <<table commands>> and 
        #store int he table_commands array
        table_commands = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    commands = re.findall(r'\<<([^>>]+)\>>', cell.text)
                    table_commands = table_commands + commands
         
        return table_commands   

        
    def generate_report(self, document_path, fields, tables): 
        """ Generates the Docx report by inserting all the job.Details fields
        in the text and inserting the results tables in the <<table command>> 
        positions
        
        Args:
            document_path (str): the path of the template doc file
            fields (str[]): Array of fields to replace with results text
            tables (ResultTable[]): The array of results tables calculated from the data
        
        """

        document = Document(document_path)

        #search through the template to find all the <<Field:___>> objects
        text_objects = []
        
        #Get all the text objects so we can replace any fields
        for paragraph in document.paragraphs:
            text_objects.append(paragraph)
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                text_objects.append(paragraph)
            for paragraph in section.footer.paragraphs:
                text_objects.append(paragraph)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text_objects.append(paragraph)
        
        #Substitute the <<Field:___>> text with the data from the fields dictionary
        for text_object in text_objects:    
             #extract the field placeholders denoated by <<command>>
             place_holders = re.findall(r'\<<Field:([^>>]+)\>>', text_object.text)
             for field in place_holders:        
                 #Date is a special field that inserts the current date
                 if field == 'Date':
                     text_object.text = text_object.text.replace('<<Field:Date>>', datetime.date.today().strftime("%d/%m/%Y"))
                 else:
                     try:
                         text_object.text = text_object.text.replace('<<Field:'+ field +'>>', fields[field])
                     except KeyError:
                         raise KeyError("Field " + field + " missing in source data");

        #Insert all the generated tables in the correct positions    
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    commands = re.findall(r'\<<([^>>]+)\>>', cell.text)            
                    for command in commands:
                        
                        #clear the placeholder
                        cell.text = ""
                        table.style = None
                        
                        if command in tables:
                            result_table = tables[command]
                            
                            if type(result_table) == list:
                                for inner_table in result_table:
                                    self.fill_table(inner_table, cell, document.styles['Table Grid'])                                  
                            else:      
                                
                                self.fill_table(result_table, cell, document.styles['Table Grid'])

        #save the document over the top of the downloaded template
        document.save(document_path)                                

    def fill_table(self, result_table, cell, style):
        """ Fills Docx tables with the actual results from the data
        
        Args:
            result_table (ResultTable): the result table with all the data
            cell (Docx.cell): The template cell that will contain the new Docx table
            with all the data.
            style (Docx.style): The formatting style for the table
        
        """
        
        data_table = result_table.table
                                
        row_count = len(data_table)
    
        shift = 0
        if result_table.title is not None:
            shift = 1
        
        column_count = len(data_table.columns)
        new_table = cell.add_table(rows=row_count+1+shift, cols=column_count)
        new_table.style = style

        if result_table.column_widths is not None:
            percents = result_table.column_widths                               

            #determine widths of asterix columns
            tot = 0
            col_defs = 0
            asterix_count = 0
            for p in percents:
                if p.isdigit():
                    tot += int(p)
                    col_defs += 1
                elif type(p) == str and p == '*':
                    asterix_count += 1
                elif type(p) == str and p == '**' and asterix_count > 0:
                    raise ValueError("Can't set column widths with ** combined with *")
            
            if asterix_count > 0:
                asterix_value = (100 - tot) / asterix_count
            else:
                asterix_value = 0
                
            #double asterix means to fill all remaining columns at the same size
            if column_count-col_defs > 0:
                double_asterix_value = (100-tot) / (column_count-col_defs)
            else:
                double_asterix_value = 0
                        
        
            for p_index in range(len(percents)):
                val = 0
                if percents[p_index].isdigit():
                    val = int(percents[p_index])
                elif type(p) == str and p == '*':
                    val = asterix_value
                elif type(p) == str and p == '**':
                    val = double_asterix_value
                new_table.columns[p_index].width = Inches(7.2 * val / 100)
   
        if result_table.title is not None:
            title_cells = new_table.rows[0].cells
            title_cell = title_cells[0].merge(title_cells[len(title_cells)-1])
            title_cell.paragraphs[0].add_run(result_table.title).bold = True
        
        hdr_cells = new_table.rows[0+shift].cells
        for column_index in range(column_count):                                
            hdr_cells[column_index].paragraphs[0].add_run(data_table.columns[column_index]).bold = True
            
        for index, df_row in data_table.iterrows():
            row_cells = new_table.rows[index+1+shift].cells                                
            for column_index in range(column_count):
                row_cells[column_index].text = df_row[data_table.columns[column_index]]
                            
    